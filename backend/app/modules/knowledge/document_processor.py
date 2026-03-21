"""Document processing for knowledge base uploads.

Supports extracting text from:
- PDF files (using PyMuPDF)
- DOCX files (using python-docx)
- TXT files (plain text)

The extracted text is then chunked and embedded for vector storage.
"""

import io
import logging
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types."""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


@dataclass
class ExtractedDocument:
    """Result of document text extraction."""

    text: str
    title: str | None
    page_count: int
    word_count: int
    document_type: DocumentType
    metadata: dict


class DocumentProcessor:
    """Process uploaded documents and extract text content."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
    MIME_TYPES = {
        "application/pdf": DocumentType.PDF,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
        "text/plain": DocumentType.TXT,
    }

    @classmethod
    def get_document_type(cls, filename: str, content_type: str | None = None) -> DocumentType:
        """Determine document type from filename or content type.

        Args:
            filename: Original filename
            content_type: MIME type if available

        Returns:
            DocumentType enum value

        Raises:
            ValueError: If document type is not supported
        """
        # Try content type first
        if content_type and content_type in cls.MIME_TYPES:
            return cls.MIME_TYPES[content_type]

        # Fall back to extension
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return DocumentType.PDF
        elif ext == ".docx":
            return DocumentType.DOCX
        elif ext == ".txt":
            return DocumentType.TXT
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {cls.SUPPORTED_EXTENSIONS}")

    @classmethod
    def extract_text(
        cls,
        file_content: bytes,
        filename: str,
        content_type: str | None = None,
    ) -> ExtractedDocument:
        """Extract text from a document.

        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type if available

        Returns:
            ExtractedDocument with extracted text and metadata

        Raises:
            ValueError: If document type is not supported
            Exception: If extraction fails
        """
        doc_type = cls.get_document_type(filename, content_type)

        if doc_type == DocumentType.PDF:
            return cls._extract_from_pdf(file_content, filename)
        elif doc_type == DocumentType.DOCX:
            return cls._extract_from_docx(file_content, filename)
        elif doc_type == DocumentType.TXT:
            return cls._extract_from_txt(file_content, filename)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

    @classmethod
    def _extract_from_pdf(cls, file_content: bytes, filename: str) -> ExtractedDocument:
        """Extract text from PDF using PyMuPDF."""
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []
            title = None

            for page_num, page in enumerate(doc):
                # Extract text from page
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)

                # Try to get title from first page metadata
                if page_num == 0 and not title:
                    title = doc.metadata.get("title")

            full_text = "\n\n".join(text_parts)

            # Clean up the text
            full_text = cls._clean_text(full_text)

            # Use filename as title if not found in metadata
            if not title:
                title = Path(filename).stem

            result = ExtractedDocument(
                text=full_text,
                title=title,
                page_count=len(doc),
                word_count=len(full_text.split()),
                document_type=DocumentType.PDF,
                metadata={
                    "author": doc.metadata.get("author"),
                    "subject": doc.metadata.get("subject"),
                    "creator": doc.metadata.get("creator"),
                    "creation_date": doc.metadata.get("creationDate"),
                },
            )

            doc.close()
            return result

        except Exception as e:
            logger.error(f"Failed to extract text from PDF {filename}: {e}")
            raise ValueError(f"Failed to process PDF: {e}") from e

    @classmethod
    def _extract_from_docx(cls, file_content: bytes, filename: str) -> ExtractedDocument:
        """Extract text from DOCX using python-docx."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            full_text = "\n\n".join(text_parts)
            full_text = cls._clean_text(full_text)

            # Get title from document properties or filename
            title = None
            if doc.core_properties.title:
                title = doc.core_properties.title
            if not title:
                title = Path(filename).stem

            return ExtractedDocument(
                text=full_text,
                title=title,
                page_count=1,  # DOCX doesn't have page concept in the same way
                word_count=len(full_text.split()),
                document_type=DocumentType.DOCX,
                metadata={
                    "author": doc.core_properties.author,
                    "subject": doc.core_properties.subject,
                    "created": str(doc.core_properties.created)
                    if doc.core_properties.created
                    else None,
                },
            )

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {filename}: {e}")
            raise ValueError(f"Failed to process DOCX: {e}") from e

    @classmethod
    def _extract_from_txt(cls, file_content: bytes, filename: str) -> ExtractedDocument:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            text = None
            for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise ValueError("Could not decode text file with any supported encoding")

            text = cls._clean_text(text)
            title = Path(filename).stem

            return ExtractedDocument(
                text=text,
                title=title,
                page_count=1,
                word_count=len(text.split()),
                document_type=DocumentType.TXT,
                metadata={},
            )

        except Exception as e:
            logger.error(f"Failed to extract text from TXT {filename}: {e}")
            raise ValueError(f"Failed to process TXT: {e}") from e

    @classmethod
    def _clean_text(cls, text: str) -> str:
        """Clean extracted text.

        - Remove excessive whitespace
        - Normalize line endings
        - Remove null characters
        """
        # Remove null characters
        text = text.replace("\x00", "")

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Remove excessive blank lines (more than 2 consecutive)
        import re

        text = re.sub(r"\n{4,}", "\n\n\n", text)

        # Remove excessive spaces
        text = re.sub(r"[ \t]{3,}", "  ", text)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text
